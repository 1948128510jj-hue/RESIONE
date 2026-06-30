"""Polish existing OKKI vs 网易外贸通 docx — improve layout without losing screenshots"""
import glob, os, copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Find the file
desktop = r'C:\Users\19481\Desktop'
files = glob.glob(desktop + r'\OKKI_vs_*.docx')
if not files:
    print("File not found!")
    exit(1)

path = files[0]
print(f"Loading: {os.path.basename(path)}")

doc = Document(path)

# ── Colors ──
ACCENT = RGBColor(0x16, 0x21, 0x3E)
GREEN = RGBColor(0x0F, 0x76, 0x6E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)
BOX_BG = RGBColor(0xEC, 0xFD, 0xF5)
RED = RGBColor(0xB9, 0x1C, 0x1C)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for attr_name, attr_val in val.items():
            element.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, before=0, after=0):
    """Set paragraph spacing"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))

# ── Process paragraph by paragraph ──
section_count = 0
modified = 0

for para in doc.paragraphs:
    text = para.text.strip()

    # Detect section titles (e.g., "1. 数据覆盖范围")
    if text and text[0].isdigit() and '. ' in text[:4] and len(text) < 40:
        # Make section title: larger, bold, colored, with top border line
        section_count += 1
        for run in para.runs:
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = ACCENT
            run.font.name = 'Arial'
        # Add bottom border spacing
        set_paragraph_spacing(para, before=240, after=120)
        modified += 1

    # Detect summary lines (starting with "总结：")
    elif text.startswith('总结：') or text.startswith('总结:'):
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            if text.startswith('总结：') or text.startswith('总结:'):
                # Make "总结：" green and bold
                if '总结' in run.text:
                    run.font.bold = True
                    run.font.color.rgb = GREEN
        set_paragraph_spacing(para, before=60, after=60)
        modified += 1

# ── Process tables ──
for i, table in enumerate(doc.tables):
    rows = table.rows
    if len(rows) == 0: continue

    # First row = header row?
    first_cell_text = rows[0].cells[0].text.strip() if rows[0].cells else ''

    # Color header rows
    for cell in rows[0].cells:
        set_cell_shading(cell, '16213E')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'

    # Alternating row colors for data rows
    for row_idx in range(1, len(rows)):
        if row_idx % 2 == 0:
            for cell in rows[row_idx].cells:
                set_cell_shading(cell, 'F1F5F9')

    # Clean up cell formatting
    for row in rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size and run.font.size < Pt(9):
                        run.font.size = Pt(10)
                    run.font.name = 'Arial'
            # Add cell padding via margins
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:w'), '60')
                el.set(qn('w:type'), 'dxa')
                tcMar.append(el)
            for side in ['left', 'right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:w'), '100')
                el.set(qn('w:type'), 'dxa')
                tcMar.append(el)
            tcPr.append(tcMar)

    modified += 1

# ── Center align cover page elements ──
for i, para in enumerate(doc.paragraphs[:10]):
    if para.text.strip():
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            if 'OKKI vs' in para.text:
                run.font.size = Pt(28)
                run.font.bold = True
            elif '对比分析框架' in para.text:
                run.font.size = Pt(18)
            elif 'RESIONE' in para.text:
                run.font.size = Pt(13)

# ── Save ──
out = desktop + r'\OKKI_vs_网易外贸通_对比框架_排版优化.docx'
doc.save(out)
print(f"Done! {section_count} sections, {modified} elements improved")
print(f"Saved: {out}")
